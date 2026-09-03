#!/usr/bin/env python3
"""Evidence-first application materials, local tracking, and reporting."""
from __future__ import annotations

import argparse
import hashlib
import html
import json
import os
import sys
import re
import shutil
import subprocess
from datetime import datetime, timedelta, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
SETTINGS = json.loads((ROOT / "config" / "settings.local.json").read_text(encoding="utf-8")) if (ROOT / "config" / "settings.local.json").exists() else {}
SIGNATURE = SETTINGS.get("candidate", {}).get("name", "<NAME>")
PYTHON = Path(os.environ.get("CAREER_AGENT_PYTHON", sys.executable))
SOFFICE = Path(os.environ.get("SOFFICE_PATH", "soffice"))


def now() -> str:
    return datetime.now(timezone.utc).isoformat()


def read_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def write_json(path: Path, value: dict) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(value, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def digest(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def job_dir(job_id: str) -> Path:
    result = ROOT / "jobs" / "inbox" / job_id
    if not (result / "job.md").exists() or not (result / "evaluation.json").exists():
        raise SystemExit(f"Job '{job_id}' must first be imported and evaluated.")
    return result


def application_dir(job_id: str) -> Path:
    return ROOT / "applications" / job_id


def select_base(job_id: str, evaluation: dict) -> str:
    text = (job_id + " " + " ".join(evaluation.get("matched_skills", []))).lower()
    if any(x in text for x in ("multimodal", "vision", "cv")):
        return "multimodal_cv"
    if any(x in text for x in ("agent", "tool", "workflow")):
        return "ai_agent"
    if any(x in text for x in ("llm", "training", "data", "annotation")):
        return "llm_training_data"
    return "ai_engineering"


def assert_provenance(base: str) -> list[dict]:
    path = ROOT / "resumes" / "base" / f"{base}.provenance.json"
    if not path.exists():
        raise SystemExit(f"Base resume provenance missing: {path}")
    data = read_json(path)
    if not data.get("bullets"):
        raise SystemExit("Refusing to generate a resume with no evidence-linked bullets.")
    return data["bullets"]


def md_to_docx(markdown: Path, output: Path) -> None:
    """Export a constrained resume Markdown document without interpreting claims."""
    doc = Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = Inches(0.65)
    section.left_margin = section.right_margin = Inches(0.75)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.05
    for raw in markdown.read_text(encoding="utf-8").splitlines():
        line = raw.strip()
        if not line or line.startswith("<!--"):
            continue
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(line[2:]); r.bold = True; r.font.size = Pt(15)
        elif line.startswith("## "):
            p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(7)
            r = p.add_run(line[3:].upper()); r.bold = True; r.font.size = Pt(11.5)
        elif line.startswith("### "):
            p = doc.add_paragraph(); r = p.add_run(line[4:]); r.bold = True
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def create_materials(args) -> None:
    source = job_dir(args.job_id)
    evaluation = read_json(source / "evaluation.json")
    base = args.base or select_base(args.job_id, evaluation)
    bullets = assert_provenance(base)
    target = ROOT / "resumes" / "generated" / args.job_id
    if target.exists() and not args.replace:
        raise SystemExit(f"Refusing to overwrite generated materials: {target}; use --replace after review.")
    target.mkdir(parents=True, exist_ok=True)
    base_md = ROOT / "resumes" / "base" / f"{base}.md"
    body = base_md.read_text(encoding="utf-8")
    job_text = (source / "job.md").read_text(encoding="utf-8")
    title = next((line[2:].strip() for line in job_text.splitlines() if line.startswith("# ")), args.job_id)
    preface = f"<!-- job_id: {args.job_id}; source_base: {base}; generated_at: {now()} -->\n"
    preface += f"<!-- target_role: {title}; fit_score: {evaluation.get('total')}; fit_label: {evaluation.get('label')} -->\n"
    resume = target / "resume.md"
    resume.write_text(preface + body, encoding="utf-8")
    provenance = {
        "job_id": args.job_id, "created_at": now(), "source_base": base,
        "job_snapshot": {"path": str(source / "job.md"), "sha256": digest(source / "job.md")},
        "evaluation": {"path": str(source / "evaluation.json"), "sha256": digest(source / "evaluation.json"), "score": evaluation.get("total")},
        "bullets": bullets, "status": "draft_requires_human_review",
    }
    write_json(target / "provenance.json", provenance)
    docx = target / "resume.docx"
    md_to_docx(resume, docx)
    print(f"Created evidence-linked draft -> {target.relative_to(ROOT)}")


def create_application(args) -> None:
    source = job_dir(args.job_id)
    app = application_dir(args.job_id)
    manifest = app / "application.json"
    if manifest.exists():
        raise SystemExit(f"Application exists: {manifest}")
    generated = ROOT / "resumes" / "generated" / args.job_id / "resume.md"
    if not generated.exists():
        raise SystemExit("Generate and review the tailored resume before creating an application archive.")
    app.mkdir(parents=True, exist_ok=True)
    snapshot = app / "job_snapshot.md"
    shutil.copyfile(source / "job.md", snapshot)
    data = {"job_id": args.job_id, "created_at": now(), "status": "draft", "events": [{"at": now(), "status": "draft", "source": "local"}], "materials": {"resume": str(generated), "provenance": str(generated.with_name("provenance.json"))}, "job_snapshot": str(snapshot)}
    write_json(manifest, data)
    print(f"Created local application archive -> {app.relative_to(ROOT)}")


def cover_letter(args) -> None:
    app = application_dir(args.job_id)
    manifest = app / "application.json"
    if not manifest.exists():
        raise SystemExit("Create the local application archive first.")
    data = read_json(manifest)
    provenance = read_json(Path(data["materials"]["provenance"]))
    evidence = ", ".join(item["fact_id"] for item in provenance["bullets"])
    text = f"# Cover Letter Draft — {args.job_id}\n\nDear Hiring Team,\n\nI am applying for this opportunity because its stated requirements align with the verified project evidence in my application materials. My relevant records are limited to the attached resume's source-linked work: {evidence}.\n\nI would welcome the opportunity to discuss how this evidence relates to the role's requirements.\n\nSincerely,\n{SIGNATURE}\n\n<!-- status: draft; evidence: {evidence} -->\n"
    (app / "cover_letter.md").write_text(text, encoding="utf-8")
    print(f"Created evidence-linked cover-letter draft -> {app.relative_to(ROOT) / 'cover_letter.md'}")


def interview_pack(args) -> None:
    source = job_dir(args.job_id)
    ev = read_json(source / "evaluation.json")
    skills = ev.get("matched_skills") or ["role requirements require manual review"]
    text = "# Interview Preparation — " + args.job_id + "\n\n## Verified evidence\n" + "\n".join(f"- {x}" for x in skills) + "\n\n## STAR prompts\n- Describe a project whose evidence is already in the submitted resume. State the situation, your verified responsibility, the action, and only verified results.\n\n## Questions to ask\n- Which deliverables define success during the first three months?\n- What data, evaluation process, and review process support this role?\n\n## Honesty boundary\nDo not claim metrics, production ownership, or tools absent from the resume provenance.\n"
    app = application_dir(args.job_id); app.mkdir(parents=True, exist_ok=True)
    (app / "interview_prep.md").write_text(text, encoding="utf-8")
    print(f"Created interview preparation -> {app.relative_to(ROOT) / 'interview_prep.md'}")


def outcome(args) -> None:
    manifest = application_dir(args.job_id) / "application.json"
    if not manifest.exists():
        raise SystemExit("Application archive not found.")
    if not args.confirm:
        raise SystemExit("Refusing to change application status without --confirm.")
    data = read_json(manifest)
    data["status"] = args.status
    data.setdefault("events", []).append({"at": now(), "status": args.status, "source": args.source})
    write_json(manifest, data)
    print(f"Recorded local application outcome: {args.status}")


def archive_submission(args) -> None:
    manifest = application_dir(args.job_id) / "application.json"
    if not manifest.exists(): raise SystemExit("Application archive not found.")
    if not args.confirm: raise SystemExit("Refusing to mark an application as submitted without --confirm.")
    data = read_json(manifest)
    stamp = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
    frozen = manifest.parent / "submitted" / stamp
    frozen.mkdir(parents=True, exist_ok=False)
    for key in ("resume", "provenance"):
        source = Path(data["materials"][key])
        shutil.copyfile(source, frozen / source.name)
    cover = manifest.parent / "cover_letter.md"
    if cover.exists(): shutil.copyfile(cover, frozen / cover.name)
    data["status"] = "applied"
    data["submitted_at"] = now()
    data.setdefault("events", []).append({"at": now(), "status": "applied", "source": "user_confirmed", "frozen_materials": str(frozen)})
    write_json(manifest, data)
    print(f"Archived submitted materials and recorded local status -> {frozen.relative_to(ROOT)}")


def followups(args) -> None:
    cutoff = datetime.now(timezone.utc) - timedelta(days=args.days)
    drafts = []
    for manifest in sorted((ROOT / "applications").glob("*/application.json")):
        data = read_json(manifest)
        if data.get("status") not in {"applied", "interview"}: continue
        submitted = data.get("submitted_at") or data.get("created_at")
        try: when = datetime.fromisoformat(submitted.replace("Z", "+00:00"))
        except (AttributeError, ValueError): continue
        if when > cutoff: continue
        job_id = data["job_id"]
        drafts.append({"job_id": job_id, "status": data["status"], "last_activity": submitted, "draft": f"Subject: Follow-up on {job_id}\n\nHello Hiring Team,\n\nI am following up on my application for {job_id}. I remain interested and would welcome any update on the process.\n\nBest regards,\n{SIGNATURE}"})
    output = ROOT / "applications" / "followup_proposals.json"
    write_json(output, {"created_at": now(), "quiet_for_days": args.days, "proposals": drafts})
    print(f"Created {len(drafts)} local follow-up proposals -> {output.relative_to(ROOT)}")


def report(_args) -> None:
    rows = []
    for path in sorted((ROOT / "applications").glob("*/application.json")):
        data = read_json(path)
        rows.append((data["job_id"], data.get("status", "unknown"), data.get("created_at", "")))
    table = "".join(f"<tr><td>{html.escape(a)}</td><td>{html.escape(b)}</td><td>{html.escape(c)}</td></tr>" for a,b,c in rows)
    output = ROOT / "reports" / "application-dashboard.html"; output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text(f"<!doctype html><meta charset='utf-8'><title>Career Agent — Applications</title><h1>Application dashboard</h1><p>Local source of truth; refreshed {html.escape(now())}.</p><table border='1'><tr><th>Job</th><th>Status</th><th>Created</th></tr>{table}</table>", encoding="utf-8")
    print(f"Created local dashboard -> {output.relative_to(ROOT)}")


def verify_resume(args) -> None:
    target = ROOT / "resumes" / "generated" / args.job_id
    docx, pdf = target / "resume.docx", target / "resume.pdf"
    if not docx.exists():
        raise SystemExit("Generate the tailored resume first.")
    if not SOFFICE.exists():
        raise SystemExit(f"LibreOffice not found: {SOFFICE}")
    converted = subprocess.run([str(SOFFICE), "--headless", "--convert-to", "pdf", "--outdir", str(target), str(docx)], capture_output=True, text=True, timeout=60)
    if converted.returncode or not pdf.exists():
        raise SystemExit(converted.stderr.strip() or converted.stdout.strip() or "DOCX to PDF conversion failed")
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise SystemExit("Install pypdf to run ATS verification: pip install pypdf") from exc
    reader = PdfReader(str(pdf))
    extracted = "\n".join(page.extract_text() or "" for page in reader.pages)
    master = read_json(ROOT / "profile" / "master" / "profile.json")
    email = master.get("identity", {}).get("public_header", {}).get("email", "")
    evaluation = read_json(ROOT / "jobs" / "inbox" / args.job_id / "evaluation.json")
    keywords = [str(item).lower() for item in evaluation.get("matched_skills", [])]
    lower = extracted.lower()
    checks = {"pdf_exists": pdf.exists(), "page_count": len(reader.pages), "email_present": bool(email and email.lower() in lower), "keyword_coverage": {word: word in lower for word in keywords}}
    render_prefix = target / "resume-preview"
    renderer = shutil.which("pdftoppm")
    render_warning = None
    if renderer:
        try:
            subprocess.run([renderer, "-f", "1", "-l", "1", "-png", str(pdf), str(render_prefix)], check=True, timeout=60, capture_output=True, text=True)
        except (subprocess.CalledProcessError, OSError) as exc:
            render_warning = f"PDF visual preview unavailable: {exc}"
    else:
        render_warning = "pdftoppm is unavailable; review the PDF directly in LibreOffice or a PDF reader."
    preview = render_prefix.with_name(render_prefix.name + "-1.png")
    if not preview.exists():
        try:
            import fitz
            document = fitz.open(pdf)
            page = document.load_page(0)
            page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False).save(preview)
            document.close()
            render_warning = None
        except ImportError:
            pass
    checks["passed"] = checks["page_count"] <= 2 and checks["email_present"]
    write_json(target / "ats_verification.json", {"verified_at": now(), "pdf": str(pdf), "text_characters": len(extracted), "rendered_first_page": str(preview), "render_warning": render_warning, "checks": checks})
    print(f"PDF/ATS verification {'passed' if checks['passed'] else 'needs review'} -> {target.relative_to(ROOT) / 'ats_verification.json'}")


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    sub = parser.add_subparsers(dest="command", required=True)
    p = sub.add_parser("generate-resume"); p.add_argument("job_id"); p.add_argument("--base", choices=["multimodal_cv", "llm_training_data", "ai_agent", "ai_engineering"]); p.add_argument("--replace", action="store_true"); p.set_defaults(func=create_materials)
    p = sub.add_parser("create-application"); p.add_argument("job_id"); p.set_defaults(func=create_application)
    p = sub.add_parser("cover-letter"); p.add_argument("job_id"); p.set_defaults(func=cover_letter)
    p = sub.add_parser("interview-pack"); p.add_argument("job_id"); p.set_defaults(func=interview_pack)
    p = sub.add_parser("outcome"); p.add_argument("job_id"); p.add_argument("status", choices=["draft", "applied", "interview", "offer", "rejected", "withdrawn", "no_response"]); p.add_argument("--source", default="user_confirmed"); p.add_argument("--confirm", action="store_true"); p.set_defaults(func=outcome)
    p = sub.add_parser("archive-submission"); p.add_argument("job_id"); p.add_argument("--confirm", action="store_true"); p.set_defaults(func=archive_submission)
    p = sub.add_parser("followups"); p.add_argument("--days", type=int, default=10); p.set_defaults(func=followups)
    p = sub.add_parser("verify-resume"); p.add_argument("job_id"); p.set_defaults(func=verify_resume)
    sub.add_parser("report").set_defaults(func=report)
    args = parser.parse_args(); args.func(args)


if __name__ == "__main__":
    main()
