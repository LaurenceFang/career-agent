#!/usr/bin/env python3
"""Build an evidence-first master career profile from indexed resume files."""
from __future__ import annotations

import argparse
import hashlib
import json
import re
from collections import defaultdict
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from pypdf import PdfReader

ROOT = Path(__file__).resolve().parents[1]
PROFILE = ROOT / "profile"


SETTINGS = read_json(ROOT / "config" / "settings.local.json") if (ROOT / "config" / "settings.local.json").exists() else {}
_candidate = SETTINGS.get("candidate", {})
IDENTITY = {
    "status": "user_authorized_use",
    "public_header": {
        "name": _candidate.get("name", "<NAME>"),
        "location": _candidate.get("location", "<CITY>, <COUNTRY>"),
        "phone": _candidate.get("phone", "<PHONE>"),
        "email": _candidate.get("email", "<EMAIL>"),
        "github": _candidate.get("github", "https://github.com/<you>"),
    },
    "aliases": _candidate.get("aliases", []),
}
_parse_rules = SETTINGS.get("parse_rules", {})
SCHOOL = _parse_rules.get("institution", "<YOUR UNIVERSITY>")
SCHOOL_NATIVE = _parse_rules.get("institution_native", "<院校名称>")
MAJOR = _parse_rules.get("major", "Artificial Intelligence")
MAJOR_NATIVE = _parse_rules.get("major_native", "人工智能")


def now(): return datetime.now(timezone.utc).isoformat()
def read_json(path): return json.loads(path.read_text(encoding="utf-8"))
def write_json(path, value):
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(value, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
def fingerprint(text): return hashlib.sha256(re.sub(r"\s+", " ", text).strip().encode("utf-8")).hexdigest()


def extract_docx(path):
    document = Document(path)
    chunks = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if values: chunks.append(" | ".join(values))
    return "\n".join(chunks)


def extract(path):
    suffix = path.suffix.lower()
    if suffix == ".docx": return extract_docx(path)
    if suffix == ".pdf": return "\n".join(page.extract_text() or "" for page in PdfReader(str(path)).pages)
    if suffix == ".md": return path.read_text(encoding="utf-8")
    raise ValueError(f"Unsupported extension: {suffix}")


def hits(text, phrases):
    lower = text.lower()
    return any(phrase.lower() in lower for phrase in phrases)


def evidence_record(item, text):
    return {"path": item["path"], "sha256": item["sha256"], "format": item["format"], "excerpt_hash": fingerprint(text), "parsed_at": now()}


def build_candidates(records):
    groups = defaultdict(list)
    def add(key, category, value, record): groups[(key, category, json.dumps(value, ensure_ascii=False, sort_keys=True))].append(record["evidence"])
    for record in records:
        text = record["text"]
        if hits(text, [SCHOOL, SCHOOL_NATIVE]) and hits(text, [MAJOR, MAJOR_NATIVE]):
            add("education_primary", "education", {"institution": SCHOOL, "degree": _parse_rules.get("degree", "<DEGREE>"), "start_date": _parse_rules.get("start_date", ""), "end_date": _parse_rules.get("end_date", ""), "end_date_status": "expected"}, record)
        for rule in _parse_rules.get("projects", []):
            if hits(text, rule.get("terms", [])):
                add(rule["key"], "projects", {"name": rule["name"]}, record)
        for skill in _parse_rules.get("skills", ["Python", "Git"]):
            if re.search(r"(?i)(?<![\w-])" + re.escape(skill) + r"(?![\w-])", text): add("skill_" + skill.lower().replace("-", "_"), "skills", {"name": skill}, record)
        for rule in _parse_rules.get("leadership", []):
            if hits(text, rule.get("terms", [])):
                add(rule["key"], "leadership", {"name": rule["name"]}, record)
    candidates = []
    for (fact_id, category, value_json), sources in sorted(groups.items()):
        value = json.loads(value_json)
        unique = {source["sha256"]: source for source in sources}
        status = "verified_consensus" if len(unique) >= 2 else "needs_review"
        candidates.append({"id": fact_id, "category": category, "value": value, "status": status, "evidence": list(unique.values())})
    return candidates


def parse_resumes(_args):
    index = read_json(ROOT / "sources" / "master_resume" / "index.json")
    records, failures = [], []
    raw_dir = PROFILE / "extractions" / "raw"
    for item in index["documents"]:
        path = Path(item["path"])
        try:
            text = extract(path)
            record = {"file": item, "status": "parsed" if len(text.strip()) >= 120 else "parsed_low_information", "character_count": len(text), "text": text, "evidence": evidence_record(item, text)}
            records.append(record)
            raw_dir.mkdir(parents=True, exist_ok=True)
            (raw_dir / f"{item['sha256']}.txt").write_text(text, encoding="utf-8")
        except Exception as error:
            failures.append({"path": item["path"], "sha256": item["sha256"], "reason": f"{type(error).__name__}: {error}"})
    manifest = {"generated_at": now(), "total": len(index["documents"]), "parsed": len(records), "failed": len(failures), "records": [{k: v for k, v in item.items() if k != "text"} for item in records], "failures": failures}
    write_json(PROFILE / "extractions" / "manifest.json", manifest)
    candidates = build_candidates(records)
    high_confidence = [candidate for candidate in candidates if candidate["status"] == "verified_consensus"]
    review = [candidate for candidate in candidates if candidate["status"] != "verified_consensus"]
    approved = read_json(PROFILE / "facts.json").get("projects", []) if (PROFILE / "facts.json").exists() else []
    approved_evidence = [{"path": item["source_path"], "sha256": item["sha256"], "format": "md", "approval": "user-approved profile sync"} for item in approved]
    master = {"schema_version": 1, "generated_at": now(), "principle": "Only verified_consensus or explicitly user-approved facts may feed base-resume drafts.", "identity": IDENTITY, "education": [item for item in high_confidence if item["category"] == "education"], "projects": [item for item in high_confidence if item["category"] == "projects"], "experience": [], "leadership": [item for item in high_confidence if item["category"] == "leadership"], "skills": [item for item in high_confidence if item["category"] == "skills"], "awards": [], "certificates": [], "availability": {"status": "needs_review"}, "career_constraints": {"status": "needs_review"}, "star_stories": [], "evidence_index": {item["id"]: item["evidence"] for item in high_confidence} | {"identity_public_header": [{"path": "multiple parsed resume sources", "approval": "user-authorized use"}], "obsidian_zhike_ai_algorithm": approved_evidence, "career_agent_mvp": [{"path": str(ROOT / "scripts" / "career_agent.py"), "approval": "local implementation evidence"}]}, "reviewed_facts": approved}
    decisions = read_json(PROFILE / "decisions.json") if (PROFILE / "decisions.json").exists() else {}
    master["skills"].extend({"id": "skill_" + item["name"].lower(), "category": "skills", "value": {"name": item["name"], "level": item["level"]}, "status": "user_confirmed", "evidence": [{"path": "user confirmation 2026-07-27"}]} for item in decisions.get("skills", []))
    master["availability"] = decisions.get("availability", master["availability"])
    master["decisions"] = decisions
    master["evidence_index"]["github_evidence"] = decisions.get("github_evidence", [])
    write_json(PROFILE / "master" / "profile.json", master)
    queues = {
        "new_facts.yaml": {"generated_at": now(), "items": [item for item in review if item["id"] not in {"skill_numpy", "skill_pytorch"}]},
        "conflicts.yaml": {"generated_at": now(), "items": []},
        "missing_information.yaml": {"generated_at": now(), "items": []},
        "suspicious_claims.yaml": {"generated_at": now(), "items": []},
    }
    for name, payload in queues.items(): write_json(PROFILE / "review_queue" / name, payload)
    print(json.dumps({"total": len(index["documents"]), "parsed": len(records), "failed": len(failures), "verified_consensus_facts": len(high_confidence), "review_items": len(review)}, ensure_ascii=False, indent=2))


def review_profile(args):
    queue_dir = PROFILE / "review_queue"
    if not args.action:
        summary = {path.name: len(read_json(path).get("items", [])) for path in queue_dir.glob("*.yaml")}
        print(json.dumps(summary, ensure_ascii=False, indent=2)); return
    if not args.queue or not args.item_id: raise SystemExit("Use --action accept|reject|hold --queue <file> --item-id <id>")
    path = queue_dir / args.queue
    payload = read_json(path)
    found = next((item for item in payload.get("items", []) if item.get("id") == args.item_id), None)
    if not found: raise SystemExit(f"Item not found: {args.item_id}")
    found["review_status"] = args.action; found["reviewed_at"] = now()
    if args.action == "accept":
        master_path = PROFILE / "master" / "profile.json"; master = read_json(master_path)
        master["reviewed_facts"].append(found); write_json(master_path, master)
    write_json(path, payload)
    print(f"{args.action}: {args.item_id}")


def main():
    parser = argparse.ArgumentParser(description=__doc__); commands = parser.add_subparsers(dest="command", required=True)
    commands.add_parser("parse-resumes").set_defaults(func=parse_resumes)
    review = commands.add_parser("review-profile"); review.add_argument("--action", choices=["accept", "reject", "hold"]); review.add_argument("--queue"); review.add_argument("--item-id"); review.set_defaults(func=review_profile)
    args = parser.parse_args(); args.func(args)


if __name__ == "__main__": main()
