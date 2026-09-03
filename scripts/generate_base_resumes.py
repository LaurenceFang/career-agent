#!/usr/bin/env python3
"""Generate base resumes in the canonical provenance format.

Each bullet written to Markdown carries a matching provenance entry with
fact_ids + evidence_paths (schemas/provenance.schema.json), so downstream
gating (scripts/hermes_resume_gate.py) can verify every claim.

DOCX export requires the optional document extra (`pip install -r
requirements-docs.txt`); Markdown output needs nothing beyond the stdlib.
"""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
SETTINGS = json.loads((ROOT / "config" / "settings.local.json").read_text(encoding="utf-8")) if (ROOT / "config" / "settings.local.json").exists() else {}
OUT = ROOT / "resumes" / "base"

COMMON = {
    "education": SETTINGS.get("candidate", {}).get("education_line", "<UNIVERSITY> — <DEGREE>, <DATES>."),
    "header": "<NAME> | <LOCATION> | github.com/<you>",
}

# Synthetic template bases showing the canonical shape. Real bases are authored
# privately from approved facts; nothing personal ships with this mirror.
RESUMES = {
    "example_agent": {
        "title": "Example AI Agent Base Resume",
        "summary": "Synthetic template resume showing the evidence-gated base-resume format. All bullets below are placeholders in the same shape as real, fact-linked bullets.",
        "skills": "Python, Git",
        "bullets": [
            ("Example Agent Project",
             "Built a local agent workflow with tool wiring, written definitions of done and human review gates; every claim links to a fact ID.",
             "project_example"),
        ],
    },
}


def write_md(key: str, data: dict) -> None:
    lines = [
        f"# {COMMON['header']}", "", f"## {data['title']}", "",
        "## Summary", data["summary"], "", "## Education", COMMON["education"],
        "<!-- evidence: education_primary -->", "", "## Selected Evidence-Linked Projects",
    ]
    provenance = []
    for i, (heading, bullet, fact) in enumerate(data["bullets"], 1):
        bullet_id = f"{key}-bullet-{i}"
        lines += [f"### {heading}", f"- {bullet}", f"<!-- evidence: {fact}; bullet_id: {bullet_id} -->", ""]
        provenance.append({
            "bullet_id": bullet_id,
            "text": bullet,
            "fact_ids": [fact],
            "evidence_paths": ["fixtures/demo/evidence/example_agent.md" if fact.startswith("project") else "fixtures/demo/evidence/coursework.md"],
        })
    lines += ["## Skills", data["skills"], "",
              "## Review notice",
              "This is a base-resume draft. It excludes unverified numbers, production claims and unresolved authorship details."]
    OUT.mkdir(parents=True, exist_ok=True)
    (OUT / f"{key}.md").write_text("\n".join(lines) + "\n", encoding="utf-8")
    (OUT / f"{key}.provenance.json").write_text(
        json.dumps({"job_id": f"base-{key}", "bullets": provenance}, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8")
    print(f"Wrote {key}.md + {key}.provenance.json -> {OUT.relative_to(ROOT)}")


def export_docx(key: str, data: dict) -> None:
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt, RGBColor
        from docx.oxml.ns import qn
    except ImportError:
        raise SystemExit("DOCX export needs the optional document extra: pip install -r requirements-docs.txt")

    def style_run(run, size=None, bold=None, color=None):
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        if size:
            run.font.size = Pt(size)
        if bold is not None:
            run.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*color)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(p.add_run("<NAME>"), 16, True, (46, 116, 181))
    p.paragraph_format.space_after = Pt(2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(p.add_run("<LOCATION> | github.com/<you>"), 9, False, (80, 80, 80))
    p.paragraph_format.space_after = Pt(10)
    for title, body in [("PROFILE", data["summary"]), ("EDUCATION", COMMON["education"])]:
        h = doc.add_paragraph()
        style_run(h.add_run(title), 12, True, (46, 116, 181))
        h.paragraph_format.space_before = Pt(8)
        h.paragraph_format.space_after = Pt(4)
        doc.add_paragraph(body)
    h = doc.add_paragraph()
    style_run(h.add_run("SELECTED PROJECTS"), 12, True, (46, 116, 181))
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(4)
    for heading, bullet, fact in data["bullets"]:
        p = doc.add_paragraph()
        style_run(p.add_run(heading), 11, True)
        p.paragraph_format.space_after = Pt(2)
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(bullet)
        p.paragraph_format.space_after = Pt(4)
        p = doc.add_paragraph()
        style_run(p.add_run(f"Evidence: {fact}"), 8, False, (100, 100, 100))
        p.paragraph_format.space_after = Pt(4)
    h = doc.add_paragraph()
    style_run(h.add_run("SKILLS"), 12, True, (46, 116, 181))
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(4)
    doc.add_paragraph(data["skills"])
    OUT.mkdir(parents=True, exist_ok=True)
    path = OUT / f"{key}.docx"
    doc.save(path)
    print(path)


def main(argv=None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--docx", action="store_true", help="also export DOCX (requires python-docx)")
    args = parser.parse_args(argv)
    for key, data in RESUMES.items():
        write_md(key, data)
        if args.docx:
            export_docx(key, data)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
